import shutil

import streamlit as st
from pytube.request import filesize
#from pytube import YouTube
from pytubefix import YouTube
from pydub import AudioSegment
import os
import io
from openai import OpenAI
import google.generativeai as geneai
from qdrant_client.http import model
from retry import retry
from pytubefix.cli import on_progress

from pydub import AudioSegment
from youtube_transcript_api.formatters import TextFormatter
from langchain_google_genai import ChatGoogleGenerativeAI

from youtube_transcript_api import YouTubeTranscriptApi
import pandas as pd
import markdown2
from docx import Document
from fpdf import FPDF
import io
from bs4 import BeautifulSoup
def germinApiKey():
    st.warning('Please enter your Google Gemini API Key')
    "[Get GOOGLE API KEY](https://ai.google.dev/)"
    openai_api_key = st.text_input(
        "GOOGLE API KEY", key="germin_api_key", type="password")
    if  "germin_api_key" not in st.session_state:

        st.session_state['germin_api_key'] = openai_api_key
        st.info("Please add your OpenAI API key to continue.")
        st.stop()
    return  st.session_state['germin_api_key']
@st.cache_resource
def YouTubeTranscript(video_id):
        try:
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            df= pd.read_csv(f'./history/data/language.csv')
            language_codes =df['code'].tolist()



            try:
               # =pd.read_csv(f'./history/data/language.csv')
               # st.dataframe(code)
               #['en', 'de', 'fr', 'es', 'sw','ru']
                transcript = transcript_list.find_generated_transcript(language_codes)  # preferred



                language_code = transcript.language_code
            except:  # Fallback if no generated transcript is found in the supported languages
                try:
                    transcript = transcript_list.find_manually_created_transcript( ['en', 'de', 'fr', 'es', 'sw','ru'])
                    language_code = transcript.language_code
                except Exception as e:
                    st.error(e)
                    return None, None  # No suitable transcript available

            fetched_transcript = transcript.fetch()
            formatter = TextFormatter()
            formatted_transcript = formatter.format_transcript(fetched_transcript)

            return formatted_transcript, language_code

        except Exception as error:
            st.error(f"Error: {error}")  # Helpful for debugging
            return None, None


@st.cache_resource
def askQuery(Ask,transcript, model,germin_key):
        try:
            llm = ChatGoogleGenerativeAI(model=model, api_key=germin_key)
            prompt = f""" Here's a transcript:\n\n
            {transcript}\n\n\n\n. You are expert on YouTube Transcript api. You can synthesize and interpret any given transcript and 
            You  provide accurate answers to this question.:{Ask}"""
            answer = llm.invoke(prompt)
            return answer.content
        except Exception as error:
            st.error(error)

def convert_markdown_to_html(markdown_text):
    """Converts Markdown to HTML."""
    return markdown2.markdown(markdown_text)

def convert_markdown_to_pdf(markdown_text):
    """Converts Markdown text to a PDF using FPDF."""

    html = markdown2.markdown(markdown_text)  # Convert markdown to HTML for better rendering

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)  # Choose a font

    lines = html.splitlines()
    for line in lines:
        if line.strip().startswith("<h1>"):
            title = line.strip()[4:-5]  # Extract title from <h1> tags
            pdf.set_font("Arial", style="B", size=18)  # Bold title
            pdf.cell(0, 10, txt=title, ln=1, align="C")  # Centered title
            pdf.set_font("Arial", size=12)  # Reset font
        elif line.strip().startswith("<h2>"):
            subtitle = line.strip()[4:-5]  # Extract title from <h1> tags
            pdf.set_font("Arial", style="B", size=16)  # Bold title
            pdf.cell(0, 10, txt=subtitle, ln=1, align="C")  # Centered title
            pdf.set_font("Arial", size=12)  # Reset font
        elif line.strip().startswith("<h3>"):
            subtitle = line.strip()[4:-5]  # Extract title from <h1> tags
            pdf.set_font("Arial", style="B", size=14)  # Bold title
            pdf.cell(0, 10, txt=subtitle, ln=1, align="C")  # Centered title
            pdf.set_font("Arial", size=12)
        elif line.strip().startswith("<p>"):
            paragraph = ((line.strip()[3:-4]).replace("<strong>", "")).replace("</strong>", "")
            paragraph = (paragraph.replace("<a href=", "")).replace('"', '')
            paragraph = paragraph.replace("</a>", "")

            # link, link_text = extract_link_and_text(paragraph)
            # if link :
            #     pdf.set_link(link)  # Set the link destination
            #     pdf.write(5, link_text)
            # else:
            pdf.multi_cell(0, 10, txt=paragraph)

            # if paragraph.strip().startswith("<strong>"):
            #     text = line.strip()[8:-9]
            #     pdf.set_font("Arial", style="B", size=12)
            #     pdf.cell(0, 10, txt=text)
            #     pdf.set_font("Arial", size=12)
            #
            # else:
            # pdf.multi_cell(0, 10, txt=paragraph)  # handle multiple lines in paragraph
        # elif line.strip().startswith("<ul>"):
        #
        #     for item in line.strip()[4:-5].split("<li>"):
        #         if item.strip():
        #             pdf.cell(10, 10, txt=". " + item.strip().replace("</li>", ""), ln=1)
        elif line.strip().startswith("<li><strong>"):

            for item in line.strip()[4:-5].split("<li><strong>"):
                item = item.replace("</strong>", "")
                item = item.replace("<strong>", "")
                item = item.replace("</ul>", "")
                if item.strip():
                    # pdf.set_font("Arial", style="B", size=12)
                    item_cell = (item.strip().replace("</li>", "")).replace("</ul>", "")
                    item_cell = (item_cell.replace("<a href=", "")).replace('"', '')
                    item_cell = item_cell.replace("</a>", "")
                    pdf.set_font("Arial", style="B", size=10)
                    pdf.cell(10, 10, txt=". " + item_cell.split(">")[0], ln=1)
                    pdf.set_font("Arial", size=12)

        else:
            text = (line.strip()).replace("</ul>", "")
            text = (text.replace("<ul>", "")).replace("<ol>", "")
            text = (text.replace("</ol>", "")).replace("</p>", "")

            pdf.multi_cell(0, 10, txt=text)

    # You can add more PDF content here if needed

    return pdf.output(dest="S").encode("latin-1")  # Return PDF as bytes

def clean_html(html_content):
    """Removes HTML tags from the given HTML content using BeautifulSoup."""
    soup = BeautifulSoup(html_content, "html.parser")
    text = soup.get_text(" ", strip=True)  # strip=True removes extra whitespace.
    return text

# @st.cache_resource
def convert_html_to_docx(html_content):
    """Converts HTML to DOCX (simplified, may not handle all HTML perfectly)."""
    doc = Document()
    #  A very basic approach -  for more robust conversion, explore dedicated HTML to DOCX libraries.
    # doc.add_paragraph(clean_html(html_content))

    lines = html_content.splitlines()
    for line in lines:

        if line.strip().startswith("<h1>"):
            title = line.strip()[4:-5]  # Extract title from <h1> tags
            doc.add_heading(title)

        elif line.strip().startswith("<h2>"):
            subtitle = line.strip()[4:-5]  # Extract title from <h1> tags
            doc.add_heading(subtitle)
        elif line.strip().startswith("<h3>"):
            subtitle = line.strip()[4:-5]  # Extract title from <h1> tags
            doc.add_section(subtitle)

        elif line.strip().startswith("<p>"):
            paragraph = ((line.strip()[3:-4]).replace("<strong>", "")).replace("</strong>", "")
            paragraph = (paragraph.replace("<a href=", "")).replace('"', '')
            paragraph = paragraph.replace("</a>", "")

            # link, link_text = extract_link_and_text(paragraph)
            # if link :
            #     pdf.set_link(link)  # Set the link destination
            #     pdf.write(5, link_text)
            # else:
            doc.add_paragraph(paragraph)

            # if paragraph.strip().startswith("<strong>"):
            #     text = line.strip()[8:-9]
            #     pdf.set_font("Arial", style="B", size=12)
            #     pdf.cell(0, 10, txt=text)
            #     pdf.set_font("Arial", size=12)
            #
            # else:
            # pdf.multi_cell(0, 10, txt=paragraph)  # handle multiple lines in paragraph
        elif line.strip().startswith("<ul>"):
            main_list = doc.add_paragraph()
            main_list.style = doc.styles['List Paragraph']
            for item in line.strip()[4:-5].split("<li>"):
                if item.strip():
                    main_list.add_run(f". {item}\n")
        #             pdf.cell(10, 10, txt=". " + item.strip().replace("</li>", ""), ln=1)
        elif line.strip().startswith("<li><strong>"):

            for item in line.strip()[4:-5].split("<li><strong>"):
                item = item.replace("</strong>", "")
                item = item.replace("<strong>", "")
                item = item.replace("</ul>", "")
                main_list = doc.add_paragraph()
                main_list.style = doc.styles['List Paragraph']
                if item.strip():
                    # pdf.set_font("Arial", style="B", size=12)
                    item_cell = (item.strip().replace("</li>", "")).replace("</ul>", "")
                    item_cell = (item_cell.replace("<a href=", "")).replace('"', '')
                    item_cell = item_cell.replace("</a>", "")
                    main_list.add_run(f". {item_cell}\n")


        else:
            text = (line.strip()).replace("</ul>", "")
            text = (text.replace("<ul>", "")).replace("<ol>", "")
            text = (text.replace("</ol>", "")).replace("</p>", "")
            doc.add_paragraph(text)

    # You can add more PDF content here if needed

    return doc

            # Clean up the temporary audio file
                #os.remove(audio_path)
#main()
video_url = st.text_input("Enter YouTube Video URL:")
if video_url  not in st.session_state:
    st.session_state['video_url']=video_url


if  st.session_state['video_url']:
    try:
        yt = YouTube(st.session_state['video_url'])
        container = st.container(border=True)
        # Display video information
        container.subheader("Video Information")
        container.write(f"Title: {yt.title}")
        container.write(f"Author: {yt.author}")
        container.write(f"Views: {yt.views}")
        container.write(f"Length: {round(yt.length/60,2)} minutes")
        if yt not in st.session_state:
            st.session_state['video_id']=yt.video_id


        container1 = st.container(border=True)
        container1.subheader("Display YouTube Video")

        # Display video
        container1.video(video_url)



        # Download video
        container2 = st.container(border=True)
        container2.subheader("You can the Download Video")
        video_stream = yt.streams.get_highest_resolution()  # Get highest resolution by default
        buffer = io.BytesIO()
        video_stream.stream_to_buffer(buffer)
        buffer.seek(0)

        container2.download_button("Download Video", data=buffer.read(), file_name=f"{yt.title}.mp4", mime="video/mp4")



        # Extract MP3 audio
        container3 = st.container(border=True)
        container3.subheader("Download Audio (MP3)")
        buffer1 = io.BytesIO()
        audio_stream = yt.streams.filter(only_audio=True).first()
        audio_stream.stream_to_buffer(buffer1)
        buffer1.seek(0)
        container3.download_button("Download Audio", data=buffer1.read(), file_name=f"{yt.title}.mp3", mime="audio/mpeg")
        container4 = st.container(border=True)

        # Extract transcript
        container4.subheader("YouTube Transcript")
        long_text = ''
        try:

            container5 = st.container(border=True)
            container4.warning('Please enter your Google Gemini API Key')
            "[Get GOOGLE API KEY](https://ai.google.dev/)"
            openai_api_key = container4.text_input(
                "GOOGLE API KEY", key="germin_api_key", type="password")
            if "germin_api_key" not in st.session_state:
                st.session_state['germin_api_key'] = openai_api_key


            if openai_api_key:
                geneai.configure(api_key=st.session_state['germin_api_key'])
                os.environ['GOOGLE_API_KEY'] = st.session_state['germin_api_key']
                choice = []
                flash_vision = []
                for m in geneai.list_models():
                    if 'generateContent' in m.supported_generation_methods:
                        #st.write(m.name)
                        model_name = m.name.split("/")[1]
                        choice.append(model_name)
                        if "2.0" in str(model_name).lower() or "-exp" in model_name:
                            flash_vision.append(model_name)
                if st.session_state['video_id'] :
                    df= pd.read_csv(f'./history/data/language.csv')
                    st.write(df['code'].tolist())
                    text, language =YouTubeTranscript(st.session_state['video_id'])
                    if text:
                        container4.subheader("Entire Transcript")
                        container4.subheader("Detected language code: "+ language)


                        container4.write(text)

                        question = container5.text_input(
                            "### Ask something about Transcript:",
                            placeholder="Can you give me a short summary in German?",
                            key='question'
                        )
                        model1 = container5.radio(
                            "Choose a Model",
                            flash_vision,
                            key='model1')
                        if  question and  model1:
                            if "model1" not in st.session_state:
                                st.session_state['model1'] = model1
                            if "question" not in st.session_state:
                                st.session_state['question'] = question
                            container5.subheader("Summary created by  AI")
                            results = askQuery(Ask=question, transcript=text, model=model1, germin_key=openai_api_key)
                            container5.markdown(results)
                            html_output = convert_markdown_to_html(results)
                            # st.button(" Download")

                            if  results :
                                try:
                                    import datetime

                                    download_format = st.radio("Download as:",
                                                               ("TXT", "PDF", "HTML", "MD", "DOCX", "PPTX"))
                                    now = datetime.datetime.now()
                                    if download_format == "TXT":
                                        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
                                        st.download_button(
                                            label="Download TXT",
                                            data=results .encode("utf-8"),  # Encode to bytes for download
                                            file_name=f"{model1}_{formatted_time}_ai_article.txt",
                                            mime="text/plain",
                                        )
                                    elif download_format == "PDF":
                                        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
                                        st.download_button(
                                            label="Download PDF",
                                            data=convert_markdown_to_pdf(results),  # Encode to bytes for download
                                            file_name=f"{model1}_{formatted_time}_ai_article.pdf",
                                            mime="application/pdf",
                                        )
                                    elif download_format == "MD":
                                        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
                                        st.download_button(
                                            label="Download Markdown",
                                            data=results.encode("utf-8"),  # Encode to bytes for download
                                            file_name=f"{model1}_{formatted_time}_ai_article.md",
                                            mime="text/plain",
                                        )
                                    # convert_markdown_to_pptx
                                    # Customize format as needed.
                                    elif download_format == "HTML":
                                        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
                                        st.download_button(
                                            label="Download HTML",
                                            data=html_output.encode("utf-8"),  # Encode to bytes for download
                                            file_name=f"{model}_{formatted_time}_ai_article.html",
                                            mime="text/html",
                                        )
                                    elif download_format == "DOCX":
                                        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
                                        docx_output = convert_html_to_docx(html_output)
                                        # Save to in-memory buffer for download
                                        with io.BytesIO() as buffer:
                                            docx_output.save(buffer)
                                            st.download_button(
                                                label="Download DOCX",
                                                data=buffer.getvalue(),
                                                file_name=f"{model}_{formatted_time}_ai_article.docx",
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" )
                                except Exception as error:
                                    st.error(error)


        except Exception as e:
            st.error(f"Error getting transcript: {e}")  # Handle potential transcript errors

    except Exception as e:
        st.error(f"Error processing video: {e}")  # Handle invalid URL or other PyTube errors
